from __future__ import annotations

import io
import os

from starlette.concurrency import run_in_threadpool

from app.core.config import get_settings
from app.modules.documents.domain.ocr import OcrEngine
from app.modules.documents.domain.parser import DocumentParser, ParsedDocument

# Extensions handled by a dedicated binary parser. Anything else (and unknown
# binary) falls back to UTF-8 text decoding.
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_XLSX_EXTS = {".xlsx", ".xlsm"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}

_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}
_XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}


def _ext(file_name: str | None) -> str:
    return os.path.splitext(file_name or "")[1].lower()


def _is_image(ext: str, mime: str) -> bool:
    return ext in _IMAGE_EXTS or mime.startswith("image/")


class MultiFormatDocumentParser(DocumentParser):
    """Extracts text from PDF, DOCX, XLSX, plain-text and image files.

    The underlying libraries (pypdf, python-docx, openpyxl) are synchronous and
    CPU-bound, so text extraction is offloaded to a worker thread. When an OCR
    engine is supplied, image files are OCR'd directly and scanned PDFs (those
    whose text layer is empty/near-empty) fall back to OCR automatically.
    """

    def __init__(self, ocr: OcrEngine | None = None) -> None:
        self._ocr = ocr
        self._pdf_min_chars = get_settings().ocr_pdf_min_chars

    def supports(self, *, file_name: str | None, mime_type: str | None) -> bool:
        return True  # plain text is the universal fallback

    async def parse(
        self, *, data: bytes, file_name: str | None, mime_type: str | None
    ) -> ParsedDocument:
        ext = _ext(file_name)
        mime = (mime_type or "").lower()

        if _is_image(ext, mime):
            return await self._parse_image(data)
        if ext in _PDF_EXTS or mime in _PDF_MIMES:
            return await self._parse_pdf(data)
        if ext in _DOCX_EXTS or mime in _DOCX_MIMES:
            return await run_in_threadpool(self._parse_docx, data)
        if ext in _XLSX_EXTS or mime in _XLSX_MIMES:
            return await run_in_threadpool(self._parse_xlsx, data)
        return await run_in_threadpool(self._parse_text, data)

    # --- image (OCR-only) ----------------------------------------------

    async def _parse_image(self, data: bytes) -> ParsedDocument:
        if self._ocr is None:
            raise RuntimeError("OCR engine is required to parse image files")
        text = await self._ocr.image_to_text(data)
        return ParsedDocument(text=text, metadata={"format": "image", "ocr": True})

    # --- pdf (text layer, OCR fallback) --------------------------------

    async def _parse_pdf(self, data: bytes) -> ParsedDocument:
        result = await run_in_threadpool(self._parse_pdf_textlayer, data)
        # A scanned PDF has pages but (almost) no extractable text — OCR it.
        if len(result.text.strip()) < self._pdf_min_chars and self._ocr is not None:
            ocr_text = await self._ocr.pdf_to_text(data)
            if len(ocr_text.strip()) > len(result.text.strip()):
                result.text = ocr_text
                result.metadata = {**result.metadata, "ocr": True}
        return result

    def _parse_pdf_textlayer(self, data: bytes) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n\n".join(p.strip() for p in pages if p.strip())
        return ParsedDocument(
            text=text,
            page_count=len(reader.pages),
            metadata={"format": "pdf", "ocr": False},
        )

    # --- office / text -------------------------------------------------

    def _parse_docx(self, data: bytes) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Pull text out of tables too — common in real-world documents.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
        text = "\n".join(parts)
        return ParsedDocument(text=text, metadata={"format": "docx"})

    def _parse_xlsx(self, data: bytes) -> ParsedDocument:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheet_names = list(wb.sheetnames)
        lines: list[str] = []
        for sheet in wb.worksheets:
            lines.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(cells):
                    lines.append("\t".join(cells))
        wb.close()
        return ParsedDocument(
            text="\n".join(lines),
            page_count=len(sheet_names),
            metadata={"format": "xlsx", "sheets": sheet_names},
        )

    def _parse_text(self, data: bytes) -> ParsedDocument:
        text = data.decode("utf-8", errors="replace")
        return ParsedDocument(text=text, metadata={"format": "text"})
